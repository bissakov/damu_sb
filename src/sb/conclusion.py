import logging
import os
import re
from datetime import datetime

from docx import Document
from docx.document import Document as DocumentObj
from docx.shared import Pt

from sb.crm import Activity, Participant
from sb.kompra import Company
from utils.utils import prettify_number

logger = logging.getLogger("DAMU")


def get_participant_list(participants: list[Participant]) -> str:
    s = ""
    for p in participants:
        if p.is_too or not p.iin:
            continue
        line = f"- {p.name}, ИИН {p.iin}\n"
        if line in s:
            continue
        s += line
    s = s.strip()
    return s


def get_guarant_list(participants: list[Participant]) -> str:
    s = ""
    for p in participants:
        if "арант" not in p.role.lower():
            continue
        line = f"- {p.name}, ИИН {p.iin}\n"
        if line in s:
            continue
        s += line
    s = s.strip()
    return s


def set_global_style(
    doc: DocumentObj, font: str = "Times New Roman", font_size: int = 12
) -> None:
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = font
            run.font.size = Pt(font_size)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = font
                        run.font.size = Pt(font_size)


def fill_conclusion_too(
    template_path: str,
    company: Company,
    owner: Participant,
    activity: Activity,
    participant_list: str,
    guarant_list: str,
) -> None:
    assert activity.guarantee

    today = datetime.fromisoformat(os.environ["today"])

    today_str = today.strftime("%d.%m.%Y")
    register_date_str = activity.guarantee.registration_date.strftime(
        "%d.%m.%Y"
    )
    if company.register_date:
        company_register_date_formatted = company.register_date.strftime(
            "%d.%m.%Y"
        )
    else:
        company_register_date_formatted = ""

    if company.last_register_date:
        if company_register_date_formatted:
            company_register_date_formatted += ", "
        company_register_date_formatted += company.last_register_date.strftime(
            "%d.%m.%Y"
        )

    doc = Document(template_path)

    paras = doc.paragraphs
    paras[8].text = paras[8].text.replace("{TOO_NAME}", company.name)
    paras[10].text = (
        paras[10]
        .text.replace("{TODAY}", today_str)
        .replace("{GUARNATEE_ID}", activity.guarantee_id)
        .replace("{REGISTER_DATE}", register_date_str)
        .replace("{TOO_NAME}", company.name)
    )
    paras[12].text = paras[12].text.replace("{TOO_NAME}", company.name)

    table = doc.tables[0]
    table.cell(0, 1).text = table.cell(0, 1).text.replace(
        "{TOO_FULLNAME}", company.full_name
    )
    table.cell(1, 1).text = table.cell(1, 1).text.replace(
        "{REGISTER_DATE}", company_register_date_formatted
    )
    table.cell(2, 1).text = table.cell(2, 1).text.replace(
        "{LAW_ADDRESS}", company.law_address
    )
    table.cell(3, 1).text = table.cell(3, 1).text.replace(
        "{BIN}", company.identifier
    )
    table.cell(4, 1).text = table.cell(4, 1).text.replace(
        "{OKED_NAME}", company.oked.name_ru
    )
    table.cell(5, 1).text = table.cell(5, 1).text.replace("{OWNER}", owner.name)
    table.cell(6, 1).text = table.cell(6, 1).text.replace(
        "{OWNER_ID}", owner.id_number or ""
    )
    table.cell(7, 1).text = table.cell(7, 1).text.replace(
        "{OWNER_IIN}", owner.iin
    )
    table.cell(8, 1).text = table.cell(8, 1).text.replace(
        "{PARTICIPANTS}", participant_list
    )
    table.cell(9, 1).text = (
        table.cell(9, 1)
        .text.replace(
            "{CREDIT_AMOUNT}", prettify_number(activity.guarantee.credit_amount)
        )
        .replace("{CREDIT_PERIOD}", str(activity.guarantee.credit_period))
    )
    table.cell(10, 1).text = table.cell(10, 1).text.replace(
        "{BANK}", activity.guarantee.bank
    )
    table.cell(11, 1).text = table.cell(11, 1).text.replace(
        "{CREDITING_PURPOSE}", activity.guarantee.crediting_purpose
    )
    table.cell(12, 1).text = (
        table.cell(12, 1)
        .text.replace(
            "{GUARANTEE_AMOUNT}",
            prettify_number(activity.guarantee.guarantee_amount),
        )
        .replace("{GUARANTEE_PERIOD}", str(activity.guarantee.guarantee_period))
    )
    table.cell(13, 1).text = table.cell(13, 1).text.replace(
        "{GUARANTS}", guarant_list
    )

    set_global_style(doc)

    company_name = re.sub(r"[^\w\- ]", "", company.name)
    conclusion_path = (
        activity.files[0].path.parent.parent
        / f"Заключение ДБ по {company_name}.docx"
    )

    doc.save(str(conclusion_path))
    logger.info(f"Conclusion saved here {conclusion_path.as_posix()!r}")


def fill_conclusion_ip(
    template_path: str,
    enterprise: Company,
    activity: Activity,
    guarant_list: str,
    coborrowers: str | None = None,
) -> None:
    assert activity.guarantee

    today = datetime.fromisoformat(os.environ["today"])

    today_str = today.strftime("%d.%m.%Y")
    register_date_str = activity.guarantee.registration_date.strftime(
        "%d.%m.%Y"
    )
    if enterprise.register_date:
        company_register_date_formatted = enterprise.register_date.strftime(
            "%d.%m.%Y"
        )
    else:
        company_register_date_formatted = ""

    if enterprise.last_register_date:
        if company_register_date_formatted:
            company_register_date_formatted += ", "
        company_register_date_formatted += (
            enterprise.last_register_date.strftime("%d.%m.%Y")
        )

    if enterprise.register_date:
        company_register_date_formatted = enterprise.register_date.strftime(
            "%d.%m.%Y"
        )
    else:
        company_register_date_formatted = ""

    if enterprise.last_register_date:
        if company_register_date_formatted:
            company_register_date_formatted += ", "
        company_register_date_formatted += (
            enterprise.last_register_date.strftime("%d.%m.%Y")
        )

    doc = Document(template_path)

    paras = doc.paragraphs
    paras[8].text = paras[8].text.replace("{IP_NAME}", enterprise.name)
    paras[10].text = (
        paras[10]
        .text.replace("{TODAY}", today_str)
        .replace("{GUARNATEE_ID}", activity.guarantee_id)
        .replace("{REGISTER_DATE}", register_date_str)
        .replace("{IP_NAME}", enterprise.name)
    )
    paras[12].text = paras[12].text.replace("{IP_NAME}", enterprise.name)

    table = doc.tables[0]
    table.cell(0, 1).text = table.cell(0, 1).text.replace(
        "{IP_FULLNAME}", enterprise.full_name
    )
    table.cell(1, 1).text = table.cell(1, 1).text.replace(
        "{REGISTER_DATE}", company_register_date_formatted
    )
    table.cell(2, 1).text = table.cell(2, 1).text.replace(
        "{OKED_NAME}", enterprise.oked.name_ru
    )
    # table.cell(3, 1).text = table.cell(3, 1).text.replace("{OWNER}", owner.name)
    # table.cell(4, 1).text = table.cell(4, 1).text.replace(
    #     "{OWNER_ID}", owner.id_number or ""
    # )
    # table.cell(5, 1).text = table.cell(5, 1).text.replace(
    #     "{OWNER_IIN}", owner.iin
    # )
    table.cell(6, 1).text = (
        table.cell(6, 1)
        .text.replace(
            "{CREDIT_AMOUNT}", prettify_number(activity.guarantee.credit_amount)
        )
        .replace("{CREDIT_PERIOD}", str(activity.guarantee.credit_period))
    )
    table.cell(7, 1).text = table.cell(7, 1).text.replace(
        "{BANK}", activity.guarantee.bank
    )
    table.cell(8, 1).text = table.cell(8, 1).text.replace(
        "{CREDITING_PURPOSE}", activity.guarantee.crediting_purpose
    )
    table.cell(9, 1).text = (
        table.cell(9, 1)
        .text.replace(
            "{GUARANTEE_AMOUNT}",
            prettify_number(activity.guarantee.guarantee_amount),
        )
        .replace("{GUARANTEE_PERIOD}", str(activity.guarantee.guarantee_period))
    )
    table.cell(11, 1).text = table.cell(11, 1).text.replace(
        "{GUARANTS}", guarant_list
    )
    if coborrowers:
        table.cell(10, 1).text = table.cell(10, 1).text.replace(
            "{CO-BORROWERS}", coborrowers
        )
    else:
        row = table.rows[10]._element
        row.getparent().remove(row)

    set_global_style(doc)

    ip_name = re.sub(r"[^\w\- ]", "", enterprise.name)
    conclusion_path = (
        activity.files[0].path.parent.parent
        / f"Заключение ДБ по {ip_name}.docx"
    )

    doc.save(str(conclusion_path))
    logger.info(f"Conclusion saved here {conclusion_path.as_posix()!r}")
