from __future__ import annotations

import dataclasses
import json
import logging
import re
from datetime import datetime
from functools import cached_property
from json.decoder import JSONDecodeError
from pathlib import Path
from types import TracebackType
from typing import Any, NamedTuple, Type, cast, override

from docx import Document
from docx.table import Table
from pydantic import BaseModel

from utils.request_handler import RequestHandler

logger = logging.getLogger("DAMU")


def parse_table(table: Table, filter_empty: bool = False) -> list[list[str]]:
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            text = cell.text.strip().replace("\n", "")
            if filter_empty and not text:
                continue
            row_data.append(text)
        if any(row_data):
            table_data.append(row_data)
    return table_data


class Participant(NamedTuple):
    role: str
    name: str
    iin: str
    id_number: str | None
    id_date: str | None
    is_too: bool


class Activity(BaseModel):
    id: str
    guarantee_id: str
    responsible_person: str
    guarantee: Guarantee | None = None
    files: list[GuaranteeFile] = dataclasses.field(default_factory=list)
    participants: list[Participant] = dataclasses.field(default_factory=list)


class Guarantee(BaseModel):
    guarantee_id: str
    bank: str
    credit_period: int
    crediting_purpose: str
    credit_amount: float
    registration_date: datetime
    guarantee_amount: float
    guarantee_period: int


@dataclasses.dataclass
class GuaranteeFile:
    id: str
    path: Path
    created_on: datetime
    type: str

    @cached_property
    def is_26(self) -> bool:
        if not self.path.name.endswith("docx"):
            return False

        docx = Document(str(self.path))
        paras = docx.paragraphs

        first_lines = " ".join(
            text
            for i in range(len(paras))
            if (text := paras[i].text.strip().lower())
        )
        first_lines = re.sub(r"[^\w \n]", "", first_lines)
        is_26 = "лица уч" in first_lines

        return is_26

    def get_participants(self) -> list[Participant]:
        if not self.is_26:
            return []

        docx = Document(str(self.path))

        table_data = parse_table(docx.tables[0])

        participants: list[Participant] = []
        for row in table_data[1:]:
            row = row[::-1]
            id_date = row[0]
            id_number = row[1] or None
            iin = row[2]
            name = row[3]
            role = row[4]
            is_too = (
                "тоо" in name.lower() or "товарищество с огр" in name.lower()
            )

            participant = Participant(
                role=role,
                name=name,
                iin=iin,
                id_number=id_number,
                id_date=id_date,
                is_too=is_too,
            )
            participants.append(participant)

        return participants


class Schemas:
    def __init__(self, schema_json_path: Path) -> None:
        self.schema_json_path = schema_json_path

        self.schemas: dict[str, Any]
        try:
            with open(schema_json_path, "r", encoding="utf-8") as f:
                self.schemas = json.load(f)
        except JSONDecodeError:
            with open(schema_json_path, "r", encoding="utf-8-sig") as f:
                self.schemas = json.load(f)

    def activities(self) -> dict[str, Any]:
        schema = self.schemas["activities"]
        if __debug__:
            schema["filters"]["items"]["6e11d8a7-f1ce-434f-8bf7-2ba016568326"][
                "items"
            ]["CustomFilters"]["items"][
                "09873e0e-f308-4681-8a30-20752ef19cfd"
            ] = {
                "filterType": 4,
                "comparisonType": 4,
                "isEnabled": True,
                "trimDateTimeParameterToDate": False,
                "leftExpression": {"expressionType": 0, "columnPath": "Status"},
                "rightExpressions": [
                    {
                        "expressionType": 2,
                        "parameter": {
                            "dataValueType": 10,
                            "value": "201cfba8-58e6-df11-971b-001d60e938c6",
                        },
                    }
                ],
            }
        return schema

    def guarantee(self, guarantee_id: str) -> dict[str, Any]:
        schema = self.schemas["guarantee"]
        schema["filters"]["items"]["primaryColumnFilter"]["rightExpression"][
            "parameter"
        ]["value"] = guarantee_id
        return schema

    def guarantee_file(self, guarantee_id: str) -> dict[str, Any]:
        schema = self.schemas["guarantee_file"]
        schema["filters"]["items"]["entityFilterGroup"]["items"][
            "masterRecordFilter"
        ]["rightExpression"]["parameter"]["value"] = guarantee_id

        schema["filters"]["items"]["entityFilterGroup"]["items"][
            "0c40db70-f3e2-4fd2-a999-e7fa53fe60cf"
        ]["rightExpression"]["parameter"]["value"] = guarantee_id

        return schema


class CRM(RequestHandler):
    def __init__(
        self,
        user: str,
        password: str,
        base_url: str,
        download_folder: Path,
        user_agent: str,
        schema_json_path: Path,
    ) -> None:
        super().__init__(user, password, base_url, download_folder)
        self.client.headers = {
            "accept": "application/json",
            "accept-language": "en-US,en;q=0.9",
            "content-type": "application/json",
            "origin": "https://crm.fund.kz",
            "priority": "u=1, i",
            "referer": "https://crm.fund.kz/Login/NuiLogin.aspx?ReturnUrl=%2f%3fsimpleLogin&simpleLogin",
            "sec-ch-ua": '"Google Chrome";v="131", "Chromium";v="131", "Not_A Brand";v="24"',
            "sec-ch-ua-mobile": "?0",
            "sec-ch-ua-platform": '"Windows"',
            "sec-fetch-dest": "empty",
            "sec-fetch-mode": "cors",
            "sec-fetch-site": "same-origin",
            "user-agent": user_agent,
            "x-request-source": "ajax-provider",
            "x-requested-with": "XMLHttpRequest",
        }

        self.schemas = Schemas(schema_json_path)
        self.is_logged_in = False

    def login(self) -> bool:
        credentials = {
            "UserName": self.user,
            "UserPassword": self.password,
            "TimeZoneOffset": -300,
        }

        logger.info("Fetching '.ASPXAUTH', 'BPMCSRF', and 'UserName' cookies")
        if not self.request(
            method="post",
            path="servicemodel/authservice.svc/login",
            json=credentials,  # type: ignore
            update_cookies=True,
        ):
            logger.error(
                "Request failed while fetching '.ASPXAUTH', 'BPMCSRF', and 'UserName' cookies"
            )
            self.is_logged_in = False
            return False
        logger.info(
            "Fetched '.ASPXAUTH', 'BPMCSRF', and 'UserName' cookies successfully"
        )

        logger.debug("Extracting 'BPMCSRF' token from cookies")
        self.client.headers["BPMCSRF"] = (
            self.client.cookies.get("BPMCSRF") or ""
        )
        logger.info("'BPMCSRF' token added to headers")

        logger.info("Login process completed successfully")
        self.is_logged_in = True
        return True

    def get_unfinished_activities(self) -> list[Activity] | None:
        if not self.is_logged_in:
            self.login()

        json_data = self.schemas.activities()

        response = self.request(
            method="post",
            path="0/DataService/json/SyncReply/SelectQuery",
            json=json_data,
        )
        if not response:
            self.is_logged_in = False
            return None

        if not hasattr(response, "json"):
            return None

        data = response.json()
        rows: list[dict[str, Any]] = data.get("rows", [])

        activities: list[Activity] = []
        for row in rows:
            guarantee = row.get("Guarantee", {})
            id = guarantee.get("value", "")
            if not id:
                continue

            activity = Activity(
                id=id,
                guarantee_id=guarantee.get("displayValue"),
                responsible_person=row.get("Owner", {}).get("displayValue"),
            )
            activities.append(activity)
        return activities

    def get_guarantee(self, guarantee_id: str) -> Guarantee | None:
        if not self.is_logged_in:
            self.login()

        json_data = self.schemas.guarantee(guarantee_id)

        response = self.request(
            method="post",
            path="0/DataService/json/SyncReply/SelectQuery",
            json=json_data,
        )
        if not response:
            self.is_logged_in = False
            return None

        if not hasattr(response, "json"):
            return None

        data = response.json()
        row: dict[str, Any] = rows[0] if (rows := data.get("rows", [])) else {}
        # bank = row.get("Bank", {}).get("displayValue")
        # credit_period = row.get("CreditPeriod")
        # crediting_purpose = row.get("CreditingPurpose", {}).get("displayValue")
        # credit_amount = row.get("CreditAmount")
        # registration_date_str = row.get("RegistrationDate")
        #
        # try:
        #     registration_date = datetime.fromisoformat(registration_date_str)
        # except ValueError as e:
        #     # FIXME TEMP
        #     raise e

        input_data = {
            "guarantee_id": guarantee_id,
            "bank": row.get("Bank", {}).get("displayValue"),
            "credit_period": row.get("CreditPeriod"),
            "crediting_purpose": row.get("CreditingPurpose", {}).get(
                "displayValue"
            ),
            "credit_amount": row.get("CreditAmount"),
            "registration_date": datetime.fromisoformat(
                row.get("RegistrationDate", "")
            ),
            "guarantee_amount": row.get("GuaranteeAmount"),
            "guarantee_period": row.get("GuaranteePeriod"),
        }

        guarantee = Guarantee(**input_data)
        return guarantee

    def download_guarantee_files(
        self, guarantee_id: str, download_folder: Path
    ) -> list[GuaranteeFile]:
        if not self.is_logged_in:
            self.login()

        json_data = self.schemas.guarantee_file(guarantee_id)

        response = self.request(
            method="post",
            path="0/DataService/json/SyncReply/SelectQuery",
            json=json_data,
        )
        if not response:
            self.is_logged_in = False
            return []

        if not hasattr(response, "json"):
            return []

        data = response.json()
        rows: list[dict[str, Any]] = data.get("rows", [])

        files: list[GuaranteeFile] = []
        for row in rows:
            file_id: str = row["Id"]
            file_name: str = row["Name"]

            # if not file_name.endswith(".docx"):
            #     continue

            created_on = datetime.fromisoformat(cast(str, row["CreatedOn"]))
            file_type: str = row["Type"]["displayValue"]

            file_name, file_ext = file_name.rsplit(".", maxsplit=1)
            file_ext = file_ext.lower()
            file_name = f"{file_name.strip()}.{file_ext}"

            file_path = download_folder / Path(
                f"{guarantee_id}/crm/{file_name}"
            )

            if not self.download_file(file_id, file_path, download_folder):
                raise Exception(
                    f"File {file_path.name!r} of {file_id!r} not downloaded"
                )

            guaranatee_file = GuaranteeFile(
                id=file_id,
                path=file_path,
                created_on=created_on,
                type=file_type,
            )

            if file_ext == "docx":
                files.append(guaranatee_file)

        return files

    def download_file(
        self, file_id: str, file_path: Path, download_folder: Path
    ) -> bool:
        if not self.is_logged_in:
            self.login()

        file_path.parent.mkdir(exist_ok=True, parents=True)

        response = self.request(
            method="get",
            path=f"0/rest/FileService/GetFile/0cf61736-d494-4029-8b1c-a33dd58edfc3/{file_id}",
        )
        if not response:
            self.is_logged_in = False
            return False

        with file_path.open("wb") as f:
            f.write(response.content)

        return True

    @override
    def __exit__(
        self,
        exc_type: Type[BaseException] | None,
        exc_val: BaseException | None,
        exc_tb: TracebackType | None,
    ) -> None:
        self.is_logged_in = False
        super().__exit__(exc_type, exc_val, exc_tb)
