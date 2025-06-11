import dataclasses
import datetime
import glob
import json
import logging
import os
import re
import sys
import zipfile
from contextlib import suppress
from datetime import date, datetime
from pathlib import Path

import docx
import dotenv
import httpx
import openpyxl
import pytz
from docx import Document
from docx.document import Document as DocumentObj
from docx.table import Table
from docx.text.paragraph import Paragraph

project_folder = Path(__file__).resolve().parent.parent.parent
os.environ["project_folder"] = str(project_folder)
os.chdir(project_folder)
sys.path.append(str(project_folder))
sys.path.append(str(project_folder / "utils"))

from sb.crm import CRM, Activity, Guarantee, GuaranteeFile
from sb.kompra import Company, Kompra, Owner
from sb.structures import Registry


def setup_logger(_today: date | None = None) -> None:
    log_format = (
        "[%(asctime)s] %(levelname)-8s %(filename)s:%(lineno)s %(message)s"
    )
    formatter = logging.Formatter(log_format, datefmt="%H:%M:%S")

    damu = logging.getLogger("DAMU")
    damu.setLevel(logging.DEBUG)

    formatter.converter = lambda *args: datetime.now(
        pytz.timezone("Asia/Almaty")
    ).timetuple()

    stream_handler = logging.StreamHandler()
    stream_handler.setLevel(logging.DEBUG)
    stream_handler.setFormatter(formatter)

    log_folder = Path("logs")
    log_folder.mkdir(exist_ok=True, parents=True)

    if _today is None:
        _today = datetime.now(pytz.timezone("Asia/Almaty")).date()

    today_str = _today.strftime("%d.%m.%y")
    year_month_folder = log_folder / _today.strftime("%Y/%B")
    year_month_folder.mkdir(parents=True, exist_ok=True)
    logger_file = year_month_folder / f"{today_str}.log"

    file_handler = logging.FileHandler(logger_file, encoding="utf-8")
    file_handler.setLevel(logging.DEBUG)
    file_handler.setFormatter(formatter)

    damu.addHandler(stream_handler)
    damu.addHandler(file_handler)


today = datetime.now(pytz.timezone("Asia/Almaty")).date()
os.environ["today"] = today.isoformat()
setup_logger(today)

logger = logging.getLogger("DAMU")


def parse_table(table: Table, filter_empty: bool = False) -> list[list[str]]:
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            text = cell.text.strip().replace("\n", "")
            if filter_empty and not text:
                continue
            row_data.append(text)
        if any(row_data):
            table_data.append(row_data)
    return table_data


def get_table(table: Table) -> list[dict[str, str]]:
    table_data = parse_table(table)
    assert all(table_data[0] == row for row in table_data)

    keys = table_data[0]
    result = [dict(reversed(list(zip(keys, row)))) for row in table_data[1:]]
    return result


@dataclasses.dataclass
class Participant:
    role: str
    name: str
    iin: str
    id_number: str | None
    id_date: date | None
    is_too: bool


def get_participant_list(participants: list[Participant]) -> str:
    s = ""
    for p in participants:
        if p.is_too or not p.iin:
            continue
        line = f"- {p.name}, ИИН {p.iin}\n"
        if line in s:
            continue
        s += line
    s = s.strip()
    return s


def get_guarant_list(participants: list[Participant]) -> str:
    s = ""
    for p in participants:
        if "арант" not in p.role.lower():
            continue
        line = f"- {p.name}, ИИН {p.iin}\n"
        if line in s:
            continue
        s += line
    s = s.strip()
    return s


def is_appendix_26(docx: DocumentObj) -> bool:
    paras = docx.paragraphs

    first_lines = " ".join(
        text
        for i in range(len(paras))
        if (text := paras[i].text.strip().lower())
    )
    first_lines = re.sub(r"[^\w \n]", "", first_lines)

    return "лица уч" in first_lines


def get_participants(file_path: Path) -> list[Participant]:
    if not file_path.name.endswith("docx"):
        return []

    docx = Document(str(file_path))
    if not is_appendix_26(docx):
        return []

    table_data = parse_table(docx.tables[0])

    participants: list[Participant] = []
    for row in table_data[1:]:
        row = row[::-1]
        try:
            id_date = datetime.strptime(row[0], "%d.%m.%Y").date()
        except ValueError:
            id_date = None
        id_number = row[1] or None
        iin = row[2]
        name = row[3]
        role = row[4]
        is_too = "тоо" in name.lower() or "товарищество с огр" in name.lower()

        participant = Participant(
            role=role,
            name=name,
            iin=iin,
            id_number=id_number,
            id_date=id_date,
            is_too=is_too,
        )
        participants.append(participant)

    return participants


def prettify_number(n: float) -> str:
    return f"{n:,.2f}".replace(",", " ").replace(".", ",")


def fill_conclusion_too(
    template_path: str,
    company: Company,
    owner: Participant,
    activity: Activity,
    participant_list: str,
    guarant_list: str,
) -> None:
    assert activity.guarantee

    today_str = today.strftime("%d.%m.%Y")
    register_date_str = activity.guarantee.registration_date.strftime(
        "%d.%m.%Y"
    )
    if company.register_date:
        company_register_date_formatted = company.register_date.strftime(
            "%d.%m.%Y"
        )
    else:
        company_register_date_formatted = ""

    if company.last_register_date:
        if company_register_date_formatted:
            company_register_date_formatted += ", "
        company_register_date_formatted += company.last_register_date.strftime(
            "%d.%m.%Y"
        )

    doc = Document(template_path)

    paras = doc.paragraphs
    paras[8].text = paras[8].text.replace("{TOO_NAME}", company.name)
    paras[10].text = (
        paras[10]
        .text.replace("{TODAY}", today_str)
        .replace("{GUARNATEE_ID}", activity.guarantee_id)
        .replace("{REGISTER_DATE}", register_date_str)
        .replace("{TOO_NAME}", company.name)
    )
    paras[12].text = paras[12].text.replace("{TOO_NAME}", company.name)

    table = doc.tables[0]
    table.cell(0, 1).text = table.cell(0, 1).text.replace(
        "{TOO_FULLNAME}", company.full_name
    )
    table.cell(1, 1).text = table.cell(1, 1).text.replace(
        "{REGISTER_DATE}", company_register_date_formatted
    )
    table.cell(2, 1).text = table.cell(2, 1).text.replace(
        "{LAW_ADDRESS}", company.law_address
    )
    table.cell(3, 1).text = table.cell(3, 1).text.replace(
        "{BIN}", company.identifier
    )
    table.cell(4, 1).text = table.cell(4, 1).text.replace(
        "{OKED_NAME}", company.oked.name_ru
    )
    table.cell(5, 1).text = table.cell(5, 1).text.replace("{OWNER}", owner.name)
    table.cell(6, 1).text = table.cell(6, 1).text.replace(
        "{OWNER_ID}", owner.id_number or ""
    )
    table.cell(7, 1).text = table.cell(7, 1).text.replace(
        "{OWNER_IIN}", owner.iin
    )
    table.cell(8, 1).text = table.cell(8, 1).text.replace(
        "{PARTICIPANTS}", participant_list
    )
    table.cell(9, 1).text = (
        table.cell(9, 1)
        .text.replace(
            "{CREDIT_AMOUNT}", prettify_number(activity.guarantee.credit_amount)
        )
        .replace("{CREDIT_PERIOD}", str(activity.guarantee.credit_period))
    )
    table.cell(10, 1).text = table.cell(10, 1).text.replace(
        "{BANK}", activity.guarantee.bank
    )
    table.cell(11, 1).text = table.cell(11, 1).text.replace(
        "{CREDITING_PURPOSE}", activity.guarantee.crediting_purpose
    )
    table.cell(12, 1).text = (
        table.cell(12, 1)
        .text.replace(
            "{GUARANTEE_AMOUNT}",
            prettify_number(activity.guarantee.guarantee_amount),
        )
        .replace("{GUARANTEE_PERIOD}", str(activity.guarantee.guarantee_period))
    )
    table.cell(13, 1).text = table.cell(13, 1).text.replace(
        "{GUARANTS}", guarant_list
    )

    company_name = re.sub(r"[^\w\- ]", "", company.name)
    conclusion_path = (
        activity.files[0].file_path.parent.parent
        / f"Заключение ДБ по {company_name}.docx"
    )

    doc.save(str(conclusion_path))
    logger.info(f"Conclusion saved here {conclusion_path.as_posix()!r}")


def fill_conclusion_ip(
    template_path: str, enterprise: Company, activity: Activity
) -> None:
    assert activity.guarantee

    today_str = today.strftime("%d.%m.%Y")
    register_date_str = activity.guarantee.registration_date.strftime(
        "%d.%m.%Y"
    )
    if enterprise.register_date:
        company_register_date_formatted = enterprise.register_date.strftime(
            "%d.%m.%Y"
        )
    else:
        company_register_date_formatted = ""

    if enterprise.last_register_date:
        if company_register_date_formatted:
            company_register_date_formatted += ", "
        company_register_date_formatted += (
            enterprise.last_register_date.strftime("%d.%m.%Y")
        )

    if enterprise.register_date:
        company_register_date_formatted = enterprise.register_date.strftime(
            "%d.%m.%Y"
        )
    else:
        company_register_date_formatted = ""

    if enterprise.last_register_date:
        if company_register_date_formatted:
            company_register_date_formatted += ", "
        company_register_date_formatted += (
            enterprise.last_register_date.strftime("%d.%m.%Y")
        )

    doc = Document(template_path)

    paras = doc.paragraphs
    paras[8].text = paras[8].text.replace("{IP_NAME}", enterprise.name)
    paras[10].text = (
        paras[10]
        .text.replace("{TODAY}", today_str)
        .replace("{GUARNATEE_ID}", activity.guarantee_id)
        .replace("{REGISTER_DATE}", register_date_str)
        .replace("{IP_NAME}", enterprise.name)
    )
    paras[12].text = paras[12].text.replace("{IP_NAME}", enterprise.name)

    table = doc.tables[0]
    table.cell(0, 1).text = table.cell(0, 1).text.replace(
        "{IP_FULLNAME}", enterprise.full_name
    )
    table.cell(1, 1).text = table.cell(1, 1).text.replace(
        "{REGISTER_DATE}", company_register_date_formatted
    )

    ip_name = re.sub(r"[^\w\- ]", "", enterprise.name)
    conclusion_path = (
        activity.files[0].file_path.parent.parent
        / f"Заключение ДБ по {ip_name}.docx"
    )

    doc.save(str(conclusion_path))
    logger.info(f"Conclusion saved here {conclusion_path.as_posix()!r}")


def main() -> None:
    logger.info("START process")

    dotenv.load_dotenv(".env")

    kompra_user = os.environ["KOMPRA_USERNAME"]
    kompra_password = os.environ["KOMPRA_PASSWORD"]
    kompra_base_url = os.environ["KOMPRA_BASE_URL"]

    registry = Registry(download_folder=Path(f"downloads/{today}"))

    crm = CRM(
        user=os.environ["CRM_USERNAME"],
        password=os.environ["CRM_PASSWORD"],
        base_url=os.environ["CRM_BASE_URL"],
        download_folder=registry.download_folder,
        user_agent=os.environ["USER_AGENT"],
        schema_json_path=registry.schema_json_path,
    )

    kompra = Kompra(
        user=kompra_user,
        password=kompra_password,
        base_url=kompra_base_url,
        download_folder=registry.download_folder,
        token_cache_path=registry.token_cache_path,
        user_agent=os.environ["USER_AGENT"],
    )

    logger.info("Data loaded")

    # with crm:
    #     activities = crm.get_unfinished_activities()
    #
    #     if not activities:
    #         logger.info(f"0 activities found...")
    #         return
    #
    #     logger.info(f"Found {len(activities)} unfinished activities...")
    #
    #     for activity in activities:
    #         logger.info(
    #             f"Working on {activity.id=!r} of {activity.responsible_person!r}"
    #         )
    #         activity.guarantee = crm.get_guarantee(activity.id)
    #         logger.info(f"Guarantee data fetched - {activity.guarantee=!r}")
    #         activity.files = crm.download_guarantee_files(
    #             activity.id, registry.download_folder
    #         )
    #         logger.info(f"Found {len(activity.files)} attached files")
    #
    #         for file in activity.files:
    #             if file.file_path.exists():
    #                 logger.info(f"{file.file_path.name} downloaded")
    #             else:
    #                 logger.error(f"{file.file_path.name} was not downloaded")
    #
    # print(activities)

    from sb.test_data import activities

    with kompra:
        for activity in activities:
            assert activity.guarantee
            for file in activity.files:
                participants = get_participants(file.file_path)
                if not participants:
                    continue

                logger.info(f"Working on {file.file_path.as_posix()!r}")

                print(participants)

                too = next((p for p in participants if p.is_too), None)
                if too:
                    logger.info(f"{too=!r}")

                    owner_participant = next(
                        p
                        for p in participants
                        if "руководитель" in p.role.lower()
                    )

                    bin = too.iin

                    enterprise = kompra.get_enterprise(bin)
                    logger.info(f"{enterprise=!r}")

                    owner = kompra.get_owner(bin)
                    logger.info(f"{owner=!r}")

                    risks = kompra.get_risks(iin=bin)

                    # summary = kompra.get_reliability_summary(bin)

                    if kompra.get_case_status(bin):
                        case_history = kompra.get_case_history(bin)
                        logger.info(f"{case_history=!r}")

                        if "Административные правонарушения" in risks:
                            if any(
                                (today.year - case.year) <= 3
                                for case in case_history.content
                                if case.type_id == 3
                            ):
                                pass
                            else:
                                logger.info(
                                    "No 'Административные правонарушения' in the past 3 years"
                                )
                                risks.pop(
                                    "Административные правонарушения", None
                                )

                        if any(
                            case.type_id == 2 for case in case_history.content
                        ):
                            risks["Уголовные разбирательства"] = True

                        if any(
                            (today.year - case.year) <= 3
                            for case in case_history.content
                            if case.type_id == 1
                        ):
                            risks["Гражданские разбирательства"] = True
                        else:
                            logger.info(
                                "No 'Гражданские разбирательства' in the past 3 years"
                            )

                    logger.info(f"{risks=!r}")

                    participant_list = get_participant_list(participants)
                    guarant_list = get_guarant_list(participants)

                    fill_conclusion_too(
                        template_path=str(registry.too_conclusion_template),
                        company=enterprise,
                        owner=owner_participant,
                        activity=activity,
                        participant_list=participant_list,
                        guarant_list=guarant_list,
                    )
                else:
                    continue
                    ip = participants[0]
                    logger.info(f"{too=!r}")

                    bin = ip.iin

                    enterprise = kompra.get_enterprise(bin)
                    logger.info(f"{enterprise=!r}")

                    # # certificate = kompra.get_certificate(iin)
                    # # logger.info(f"{certificate=!r}")
                    # #
                    # # properties = kompra.get_properties(iin)
                    # # logger.info(f"{properties=!r}")
                    # #
                    # # adm_fines = kompra.get_adm_fines(iin)
                    # # logger.info(f"{adm_fines=!r}")
                    # #
                    # # risk_tax_arrear = kompra.risk_tax_arrear_status(iin)
                    # # logger.info(f"{risk_tax_arrear=!r}")
                    # #
                    # # risk_mock_taxpayer = kompra.risk_mock_taxpayer(iin)
                    # # logger.info(f"{risk_mock_taxpayer=!r}")
                    #
                    # summary = kompra.get_reliability_summary(bin)
                    # logger.info(f"{summary=!r}")

                    fill_conclusion_ip(
                        template_path=str(registry.ip_conclusion_template),
                        enterprise=enterprise,
                        activity=activity,
                    )

                print("=" * 100)


if __name__ == "__main__":
    try:
        main()
    finally:
        logger.info("FINISH process")
